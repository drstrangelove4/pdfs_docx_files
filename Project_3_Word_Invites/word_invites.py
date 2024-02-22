import docx

PATH = "guests.txt"

# -----------------------------------------------------------------------------------------------------------------------


def get_names(invite_path):
    """
    Takes a path of a text file as input and returns a list of names.
    """

    # Open the text file, strip whitespace and create guest namelist.
    guest_names = []
    with open(invite_path) as file:
        guests = file.readlines()
        for name in guests:
            guest_names.append(name.rstrip())

    return guest_names


# -----------------------------------------------------------------------------------------------------------------------


def make_invite(guest_list):
    """
    Takes a list of names and creates an 'invite' for them.
    """
    for name in guest_list:
        print(f"Creating invite for {name}:")

        # Create and add content to the docx
        new_invite = docx.Document()
        new_invite.add_heading(f"Dear {name}", 2)
        new_invite.add_paragraph("Please do not come to my birthday.")
        new_invite.add_paragraph("Peace out bye.")

        # save the 'invite' to the person's name.
        new_invite.save(f"message_for_{name}.docx")
        print(f"Invite for {name} has been saved")


# -----------------------------------------------------------------------------------------------------------------------


def main():
    # Make an invite list out of the names at the text file in path. If empty raise an exception.
    invite_names = get_names(PATH)
    if invite_names:
        make_invite(guest_list=invite_names)
    else:
        raise Exception("Empty invite list.")


# -----------------------------------------------------------------------------------------------------------------------


if __name__ == "__main__":
    main()
